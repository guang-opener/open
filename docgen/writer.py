"""
Word 文档生成模块 - 基于 python-docx 的科技调研报告生成器

功能:
- 标题页 (中英文标题、日期)
- 多级标题 (Heading 1-3)
- 正文段落 (首行缩进)
- 格式化表格 (表头着色、边框)
- 项目符号列表
- 参考文献
- 自定义字体/字号/颜色
- A4 页面设置
"""

import logging
import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


# ============================================================
# 默认配置
# ============================================================

DEFAULT_STYLES = {
    "title_font": "微软雅黑",
    "heading_font": "微软雅黑",
    "body_font": "宋体",
    "title_size": 26,
    "heading1_size": 18,
    "heading2_size": 15,
    "heading3_size": 13,
    "body_size": 12,
    "title_color": "1F4E79",
    "heading_color": "1F4E79",
    "body_color": "000000",
    "page_size": "A4",
    "margin_top_cm": 2.54,
    "margin_bottom_cm": 2.54,
    "margin_left_cm": 2.5,
    "margin_right_cm": 2.5,
}


# ============================================================
# 文档生成器
# ============================================================

class ReportWriter:
    """科技调研报告 Word 文档生成器"""

    def __init__(self, styles: dict = None):
        """
        Args:
            styles: 样式覆盖配置 (与 DEFAULT_STYLES 合并)
        """
        self.styles = {**DEFAULT_STYLES, **(styles or {})}
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        """设置页面"""
        section = self.doc.sections[0]
        if self.styles["page_size"] == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        section.top_margin = Cm(self.styles["margin_top_cm"])
        section.bottom_margin = Cm(self.styles["margin_bottom_cm"])
        section.left_margin = Cm(self.styles["margin_left_cm"])
        section.right_margin = Cm(self.styles["margin_right_cm"])

    def _setup_styles(self):
        """设置默认样式"""
        style = self.doc.styles["Normal"]
        style.font.name = self.styles["body_font"]
        style.font.size = Pt(self.styles["body_size"])
        style.element.rPr.rFonts.set(qn("w:eastAsia"), self.styles["body_font"])

    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        """十六进制颜色转 RGBColor"""
        hex_color = hex_color.lstrip("#")
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )

    def _set_font(self, run, font_name: str, size: int, color: str = None, bold: bool = False):
        """设置字体属性 (支持中文字体)"""
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        # 设置中文字体
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), font_name)
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        if color:
            run.font.color.rgb = self._hex_to_rgb(color)

    def _add_paragraph_with_style(
        self,
        text: str,
        font_name: str = None,
        font_size: int = None,
        color: str = None,
        bold: bool = False,
        alignment: WD_ALIGN_PARAGRAPH = None,
        space_before: int = 6,
        space_after: int = 6,
        first_line_indent: float = None,  # 首行缩进 (字符数)
        line_spacing: float = 1.5,
    ):
        """添加格式化段落"""
        para = self.doc.add_paragraph()
        if alignment is not None:
            para.alignment = alignment

        pf = para.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = line_spacing

        if first_line_indent:
            # 首行缩进: 字号 * 字符数 = 磅值
            indent_pt = (font_size or self.styles["body_size"]) * first_line_indent
            pf.first_line_indent = Pt(indent_pt)

        run = para.add_run(text)
        fn = font_name or self.styles["body_font"]
        fs = font_size or self.styles["body_size"]
        self._set_font(run, fn, fs, color, bold)
        return para

    # ---- 标题页 ----

    def add_title_page(self, title: str, subtitle_en: str = "", subtitle_cn: str = "", date: str = ""):
        """添加标题/封面页"""
        # 空行
        for _ in range(4):
            self._add_paragraph_with_style("", space_before=12, space_after=0)

        # 主标题
        self._add_paragraph_with_style(
            title,
            font_name=self.styles["title_font"],
            font_size=self.styles["title_size"],
            color=self.styles["title_color"],
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=60,
            space_after=20,
        )

        # 英文副标题
        if subtitle_en:
            self._add_paragraph_with_style(
                subtitle_en,
                font_name="Arial",
                font_size=14,
                color="888888",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=6,
                space_after=4,
            )

        # 中文副标题
        if subtitle_cn:
            self._add_paragraph_with_style(
                subtitle_cn,
                font_name=self.styles["title_font"],
                font_size=13,
                color="999999",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=4,
                space_after=20,
            )

        # 日期
        if date:
            self._add_paragraph_with_style(
                date,
                font_name=self.styles["title_font"],
                font_size=12,
                color="AAAAAA",
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=60,
                space_after=0,
            )

        # 分页
        self.doc.add_page_break()

    # ---- 标题 ----

    def add_heading1(self, text: str):
        """一级标题"""
        self._add_paragraph_with_style(
            text,
            font_name=self.styles["heading_font"],
            font_size=self.styles["heading1_size"],
            color=self.styles["heading_color"],
            bold=True,
            space_before=24,
            space_after=12,
        )

    def add_heading2(self, text: str):
        """二级标题"""
        self._add_paragraph_with_style(
            text,
            font_name=self.styles["heading_font"],
            font_size=self.styles["heading2_size"],
            color="2E75B6",
            bold=True,
            space_before=18,
            space_after=8,
        )

    def add_heading3(self, text: str):
        """三级标题"""
        self._add_paragraph_with_style(
            text,
            font_name=self.styles["heading_font"],
            font_size=self.styles["heading3_size"],
            color="3A8FD4",
            bold=True,
            space_before=12,
            space_after=6,
        )

    # ---- 正文 ----

    def add_body(self, text: str, indent: bool = True):
        """
        添加正文段落

        Args:
            text: 段落文本
            indent: 是否首行缩进两个字符
        """
        self._add_paragraph_with_style(
            text,
            font_name=self.styles["body_font"],
            font_size=self.styles["body_size"],
            color=self.styles["body_color"],
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=4,
            space_after=4,
            first_line_indent=2.0 if indent else 0,
            line_spacing=1.5,
        )

    def add_bullet(self, text: str, indent_level: int = 0):
        """添加项目符号条目"""
        left_indent = 0.75 + indent_level * 0.5  # cm
        self._add_paragraph_with_style(
            f"• {text}",
            font_name=self.styles["body_font"],
            font_size=self.styles["body_size"],
            color=self.styles["body_color"],
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_before=2,
            space_after=2,
            first_line_indent=0,
            line_spacing=1.3,
        )
        # 设置左缩进
        para = self.doc.paragraphs[-1]
        para.paragraph_format.left_indent = Cm(left_indent)

    def add_body_from_markdown(self, markdown: str):
        """
        从 Markdown 文本自动插入格式化内容

        支持的语法:
        - # ## ### 标题
        - 普通文本为正文
        - - 开头为项目符号
        - **text** 加粗
        - 自动首行缩进
        """
        lines = markdown.strip().split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # 标题
            if line.startswith("### "):
                self.add_heading3(line[4:])
            elif line.startswith("## "):
                self.add_heading2(line[3:])
            elif line.startswith("# "):
                self.add_heading1(line[2:])
            # 项目符号
            elif line.startswith("- ") or line.startswith("* ") or line.startswith("• "):
                bullet_text = line[2:].strip()
                # 去除内联加粗标记
                bullet_text = bullet_text.replace("**", "")
                self.add_bullet(bullet_text)
            # 编号列表
            elif re.match(r"^\d+[\.\)]\s", line):
                num_text = re.sub(r"^\d+[\.\)]\s", "", line).strip()
                num_text = num_text.replace("**", "")
                self.add_bullet(num_text)
            # 分隔线
            elif line.startswith("---") or line.startswith("***"):
                pass  # 跳过分隔线
            # 表格行 (以 | 开始)
            elif line.startswith("|") and line.endswith("|"):
                pass  # 表格由 add_comparison_table 单独处理
            # 正文
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
                self.add_body(clean)

            i += 1

    # ---- 表格 ----

    def add_table(
        self,
        headers: list[str],
        rows: list[list[str]],
        col_widths: list[float] = None,
        caption: str = "",
    ):
        """
        添加格式化表格

        Args:
            headers: 表头列表
            rows: 数据行列表
            col_widths: 列宽 (厘米), None 则自动
            caption: 表格标题
        """
        if caption:
            self._add_paragraph_with_style(
                caption,
                font_name=self.styles["body_font"],
                font_size=10,
                color="666666",
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=12,
                space_after=4,
            )

        n_cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # 表头
        header_cells = table.rows[0].cells
        for j, header_text in enumerate(headers):
            cell = header_cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(header_text)
            self._set_font(run, self.styles["heading_font"], 10, "FFFFFF", bold=True)
            # 表头背景色
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{self.styles["heading_color"]}" w:val="clear"/>'
            )
            cell._element.get_or_add_tcPr().append(shading)

        # 数据行
        for i, row_data in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_text in enumerate(row_data):
                if j >= n_cols:
                    break
                cell = row_cells[j]
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(str(cell_text))
                self._set_font(run, self.styles["body_font"], 9, "000000")
                # 交替行背景
                if i % 2 == 0:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F2F7FB" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)

        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                if i < n_cols:
                    for row in table.rows:
                        row.cells[i].width = Cm(width)

        # 空行
        self._add_paragraph_with_style("", space_before=4, space_after=4)

    def add_comparison_table(self, comparison_data: dict):
        """
        从字典数据添加对比表

        comparison_data = {
            "headers": ["对比维度", "方案A", "方案B"],
            "rows": [
                ["参数1", "值", "值"],
                ...
            ]
        }
        """
        headers = comparison_data.get("headers", [])
        rows = comparison_data.get("rows", [])
        if headers and rows:
            self.add_table(headers, rows, caption="技术对比表")

    # ---- 参考文献 ----

    def add_references(self, references: list[str]):
        """添加参考文献列表"""
        self.add_heading1("参考文献")
        for i, ref in enumerate(references, 1):
            self._add_paragraph_with_style(
                f"[{i}] {ref}",
                font_name=self.styles["body_font"],
                font_size=10,
                color="555555",
                space_before=2,
                space_after=2,
                first_line_indent=0,
                line_spacing=1.2,
            )

    # ---- 图片 ----

    def add_image_placeholder(self, description: str):
        """添加图片占位符 (提示用户手动插入)"""
        self._add_paragraph_with_style(
            f"[插图位置: {description}]",
            font_name=self.styles["body_font"],
            font_size=10,
            color="CC0000",
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=12,
            space_after=4,
        )

    def add_image(self, image_path: str, width_cm: float = 14.0, caption: str = ""):
        """
        插入图片

        Args:
            image_path: 图片文件路径
            width_cm: 图片宽度 (厘米)
            caption: 图片标题
        """
        try:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Cm(width_cm))

            if caption:
                self._add_paragraph_with_style(
                    caption,
                    font_name=self.styles["body_font"],
                    font_size=9,
                    color="666666",
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=2,
                    space_after=12,
                )
        except FileNotFoundError:
            logger.warning(f"Image not found: {image_path}")
            self.add_image_placeholder(f"未找到图片: {image_path}")

    # ---- 保存 ----

    def save(self, filepath: str):
        """保存文档"""
        # 确保输出目录存在
        Path(filepath).parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(filepath)
        logger.info(f"Document saved to {filepath}")
        return filepath


# ============================================================
# 便捷函数: 从 AnalyzedReport 直接生成文档
# ============================================================

def generate_report_from_analysis(
    report_data,  # AnalyzedReport
    output_path: str,
    styles: dict = None,
    topic_en: str = "",
    topic_subtitle: str = "",
    date: str = "",
    comparison_data: dict = None,
    image_paths: list[str] = None,
) -> str:
    """
    从分析报告数据生成 Word 文档

    Args:
        report_data: AnalyzedReport 实例
        output_path: 输出路径
        styles: 样式配置
        topic_en: 英文主题
        topic_subtitle: 副标题
        date: 日期
        comparison_data: 对比表数据
        image_paths: 待插入的图片路径

    Returns:
        保存的文档路径
    """
    writer = ReportWriter(styles)

    # 标题页
    writer.add_title_page(
        title=report_data.topic,
        subtitle_en=topic_en,
        subtitle_cn=topic_subtitle,
        date=date,
    )

    # 逐节写入
    for section in report_data.sections:
        writer.add_heading1(section.title)

        if section.content:
            writer.add_body_from_markdown(section.content)

        for sub in section.subsections:
            writer.add_heading2(sub.title)
            if sub.content:
                writer.add_body_from_markdown(sub.content)

    # 对比表
    if comparison_data:
        writer.add_heading1("技术对比分析")
        writer.add_comparison_table(comparison_data)

    # 图片
    if image_paths:
        for i, img_path in enumerate(image_paths):
            writer.add_image(img_path, caption=f"图{i+1}")

    # 参考来源汇总
    refs = []
    for item in report_data.item_scores:
        if item.relevance_score >= 60:
            authors = ", ".join(item.raw_item.get("authors", [])[:3])
            year = item.raw_item.get("year", "")
            journal = item.raw_item.get("journal", "") or item.raw_item.get("source_name", "")
            url = item.raw_item.get("source_url", "")
            ref_text = f'{authors}. "{item.title}". {journal}, {year}.'
            if url:
                ref_text += f" {url}"
            refs.append(ref_text)
    writer.add_references(refs)

    return writer.save(output_path)


# ============================================================
# 简易报告生成 (不依赖 AnalyzedReport)
# ============================================================

def quick_report(
    topic: str,
    sections: dict,
    output_path: str,
    topic_en: str = "",
    subtitle: str = "",
    date: str = "",
    tables: list[dict] = None,
    references: list[str] = None,
    styles: dict = None,
) -> str:
    """
    快速生成报告 (简化 API)

    Args:
        topic: 报告标题
        sections: {章节标题: 内容文本(Markdown)} 字典
        output_path: 输出路径
        topic_en: 英文标题
        subtitle: 副标题
        date: 日期
        tables: [{headers, rows, caption}] 表格列表
        references: 参考文献列表
        styles: 样式覆盖

    Returns:
        保存的文件路径
    """
    from datetime import datetime

    writer = ReportWriter(styles)
    date = date or datetime.now().strftime("%Y年%m月")

    writer.add_title_page(topic, topic_en, subtitle, date)

    for section_title, content in sections.items():
        writer.add_heading1(section_title)
        if isinstance(content, str):
            writer.add_body_from_markdown(content)
        elif isinstance(content, list):
            for item in content:
                if isinstance(item, str):
                    writer.add_body(item)
                elif isinstance(item, dict):
                    if item.get("type") == "bullet":
                        writer.add_bullet(item["text"])
                    elif item.get("type") == "subheading":
                        writer.add_heading2(item["text"])

    if tables:
        for tbl in tables:
            writer.add_table(
                headers=tbl.get("headers", []),
                rows=tbl.get("rows", []),
                caption=tbl.get("caption", ""),
            )

    if references:
        writer.add_references(references)

    return writer.save(output_path)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

    # 测试: 生成一份快速报告
    sections = {
        "一、摘要": "本文系统调研了REBCO高温超导带材接头技术的最新进展，重点分析了钎焊、超声辅助焊接和铜扩散连接三种主流工艺路线的技术特点。研究发现，铜扩散连接技术具有最低的接头电阻（约16.8 nΩ·cm²），而超声辅助焊接在免纤剂、高可靠性方面具有显著优势。",
        "二、技术背景": """超导带材接头是超导电力设备工程化的核心关键技术之一。接头的基本要求包括：

- 低电阻：接头电阻应尽量接近超导体本征电阻
- 高机械强度：承受磁体绕制和运行过程中的机械应力
- 良好的热稳定性：耐受冷却-回温循环
- 工艺可靠性：操作简便、一致性好""",
        "三、专利分析": "通过对Google Patents和Semantic Scholar的检索分析，共筛选出高度相关专利15件、学术论文23篇。主要专利权人包括美国SuperPower公司、韩国SuNAM公司、日本藤仓株式会社和中国上海超导科技股份有限公司。",
    }

    quick_report(
        topic="超导带材接头技术调研报告",
        topic_en="Survey of Superconducting Tape Joint Technologies",
        subtitle="—— 钎焊、超声焊接与铜扩散连接技术综合评述 ——",
        sections=sections,
        output_path="test_report.docx",
    )
